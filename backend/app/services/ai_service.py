from __future__ import annotations

import json
import logging
import os
import sys
import time
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
from pathlib import Path
from typing import Any

import vertexai 
from vertexai.generative_models import GenerativeModel, Part, GenerationConfig

from ..config import settings

logger = logging.getLogger(__name__)

_PROMPT_DIR = Path(__file__).resolve().parent.parent / "Prompt"
_PROMPT_FILE = _PROMPT_DIR / "evidence_evaluation_V1.txt"
_REPORT_PROMPT_FILE = _PROMPT_DIR / "report_generation_V1.txt"
_AWS_AUTOFILL_PROMPT_FILE = _PROMPT_DIR / "aws_autofill_V1.txt"
_GCP_AUTOFILL_PROMPT_FILE = _PROMPT_DIR / "gcp_autofill_V1.txt"
_loaded_prompt_template: str | None = None
_AI_LOG_PREVIEW_CHARS = 8000
# Max chars of raw model text printed to CMD (full JSON can be huge). Override with env AI_LOG_RAW_MAX_CHARS.
_AI_RAW_MAX_CONSOLE = int(os.getenv("AI_LOG_RAW_MAX_CHARS", "500000"))

_LEVEL_TAG = {
    "INFO": "INFO ",
    "WARN": "WARN ",
    "ERROR": "ERROR",
}


def _ai_print_console(text: str, *, end: str = "\n") -> None:
    """
    Single write to **stderr** (typical server console). We do not mirror to stdout: that caused every
    `[AI]` line to appear twice. Uvicorn's own `logger.info` on stderr would triple the noise — see `_ai_emit`.
    """
    try:
        sys.stderr.write(text + end)
        sys.stderr.flush()
    except Exception:
        pass


def _ai_emit(level: str, operation: str, message: str, *, exc_info: bool = False) -> None:
    """
    One human-readable `[AI]` line on stderr. Avoid `logger.info`/`warning` here: Uvicorn's default handler
    also prints those to the same console → duplicate blocks that look like “the model ran many times”.
    Use `logger.debug` so aggregators can still enable `DEBUG` for `app.services.ai_service` without doubling output.
    """
    tag = _LEVEL_TAG.get(level.upper(), level[:5].ljust(5))
    line = f"[AI][{tag}][{operation}] {message}"
    _ai_print_console(line)
    if level.upper() == "ERROR":
        logger.error("[%s] %s", operation, message, exc_info=exc_info)
    else:
        logger.debug("[%s] %s", operation, message)


def _ai_raw_response_console(operation: str, raw_text: str) -> None:
    """Print full raw model output between markers so CMD shows the exact AI response."""
    n = len(raw_text)
    body = raw_text
    if n > _AI_RAW_MAX_CONSOLE:
        body = raw_text[:_AI_RAW_MAX_CONSOLE]
        _ai_emit(
            "WARN",
            operation,
            f"Console raw truncated: {n} chars > AI_LOG_RAW_MAX_CHARS={_AI_RAW_MAX_CONSOLE}",
        )
    _ai_print_console(f"[AI][RESPONSE][{operation}] ===== RAW BEGIN ({len(body)} of {n} chars) =====")
    _ai_print_console(body.rstrip("\r\n") + "\n")
    _ai_print_console(f"[AI][RESPONSE][{operation}] ===== RAW END =====")
    cap = 12000
    tail = "" if n <= cap else f"\n...[truncated {n - cap} chars]"
    logger.debug("[%s] Raw response (%d chars) log preview:%s%s", operation, n, tail, raw_text[:cap])


def _load_prompt_template() -> str:
    """Load the evidence evaluation prompt template (always re-read in dev for easier iteration)."""
    global _loaded_prompt_template
    if not _PROMPT_FILE.is_file():
        raise FileNotFoundError(f"Prompt file not found: {_PROMPT_FILE}")
    _loaded_prompt_template = _PROMPT_FILE.read_text(encoding="utf-8")
    return _loaded_prompt_template

_model: GenerativeModel | None = None


def _get_model() -> GenerativeModel:
    """Initialise Vertex AI using Google ADC and return the cached model.

    Uses ``GOOGLE_CLOUD_PROJECT`` and ``VERTEX_AI_LOCATION`` from settings.
    Credentials come from ``gcloud auth application-default login`` (ADC).
    """
    global _model
    if _model is None:
        project = settings.GOOGLE_CLOUD_PROJECT
        if not project:
            raise ValueError(
                "GOOGLE_CLOUD_PROJECT is not set. Add GOOGLE_CLOUD_PROJECT=<your-project-id> to backend/.env."
            )

        _ai_emit("INFO", "vertex_init", f"project={project} location={settings.VERTEX_AI_LOCATION} model={settings.VERTEX_AI_MODEL}")
        t0 = time.monotonic()
        vertexai.init(
            project=project,
            location=settings.VERTEX_AI_LOCATION,
        )
        _model = GenerativeModel(settings.VERTEX_AI_MODEL)
        _ai_emit("INFO", "vertex_init", f"ok elapsed_s={time.monotonic() - t0:.1f}")
    return _model


def prepare_file_part(file_path_or_bytes: str | bytes, mime_type: str) -> Part:
    """Return a Vertex AI Part for LLM to read. Supports PDF, images, Excel, Word, CSV, text.
    PDF/images: sent as binary (model reads natively). Excel/Word/CSV/text: extracted as text."""
    if isinstance(file_path_or_bytes, bytes):
        raw = file_path_or_bytes
    else:
        with open(file_path_or_bytes, "rb") as f:
            raw = f.read()

    mime_lower = (mime_type or "").lower()

    # PDF and images: send as binary so the model can read them directly
    if mime_lower in (
        "application/pdf",
        "image/png",
        "image/jpeg",
        "image/jpg",
        "image/webp",
        "image/gif",
    ):
        return Part.from_data(raw, mime_type=mime_lower or "application/octet-stream")

    # Excel (.xlsx, .xls): convert to CSV text
    if "spreadsheet" in mime_lower or "excel" in mime_lower or "sheet" in mime_lower:
        import io

        try:
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(raw), data_only=True)
            text_parts: list[str] = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                rows = [[str(c.value or "") for c in row] for row in ws.iter_rows()]
                csv_text = "\n".join(",".join(r) for r in rows)
                text_parts.append(f"[Sheet: {sheet}]\n{csv_text}")
            return Part.from_text("\n\n".join(text_parts))
        except Exception:
            return Part.from_text(raw.decode("utf-8", errors="replace"))

    # Word (.docx): extract text with python-docx
    if "wordprocessingml" in mime_lower or "msword" in mime_lower:
        import io

        try:
            from docx import Document

            doc = Document(io.BytesIO(raw))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [str(c.text or "").strip() for c in row.cells]
                    if any(cells):
                        tables_text.append(" | ".join(cells))
            return Part.from_text(
                "\n\n".join(paragraphs) + ("\n\n[Tables]\n" + "\n".join(tables_text) if tables_text else "")
            )
        except Exception:
            return Part.from_text(raw.decode("utf-8", errors="replace"))

    # CSV, plain text, markdown, etc.: decode as text
    return Part.from_text(raw.decode("utf-8", errors="replace"))


def _parse_numbered_criteria(value: str | dict | Any | None) -> list[tuple[str, str]]:
    """Parse sufficiency/evaluation_criteria JSON (string or dict) into [(id, label), ...], keys sorted numerically.
    Handles sufficiency_criteria as {"sufficiency_criteria": ["a","b"]} -> [(1,a),(2,b)].
    Does not handle evaluation_criteria pass_if/fail_if/cross_checks; use _parse_sufficiency_json and
    _parse_evaluation_criteria_for_prompt / _eval_criteria_pass_if_only for those."""
    if value is None:
        return []
    if isinstance(value, dict):
        obj = value
    elif isinstance(value, str):
        value = value.strip()
        if not value:
            return []
        try:
            obj = json.loads(value)
        except Exception:
            return []
    else:
        return []
    if not isinstance(obj, dict):
        return []
    # Array under key "sufficiency_criteria" (from 2025 JSON sheet)
    if "sufficiency_criteria" in obj and isinstance(obj["sufficiency_criteria"], list):
        return [(str(i), str(x).strip()) for i, x in enumerate(obj["sufficiency_criteria"], 1) if str(x).strip()]
    keys = sorted(obj.keys(), key=lambda k: (int(k) if str(k).isdigit() else 999, k))
    return [(k, str(obj[k]).strip()) for k in keys if str(obj[k]).strip()]


def _parse_evaluation_criteria_structured(ev_raw: str | dict | Any | None) -> dict | None:
    """If evaluation_criteria is {pass_if, fail_if, cross_checks}, return the dict; else None."""
    if ev_raw is None:
        return None
    if isinstance(ev_raw, dict):
        obj = ev_raw
    elif isinstance(ev_raw, str):
        ev_raw = ev_raw.strip()
        if not ev_raw:
            return None
        try:
            obj = json.loads(ev_raw)
        except Exception:
            return None
    else:
        return None
    if not isinstance(obj, dict) or "pass_if" not in obj:
        return None
    return obj


def _format_evaluation_structured_for_prompt(control_id: str, ev_obj: dict) -> str:
    """Format pass_if, fail_if, cross_checks into prompt lines with stable IDs (eval_1.., eval_f1.., eval_c1..)."""
    lines: list[str] = []
    idx = 0
    for label in ev_obj.get("pass_if") or []:
        idx += 1
        lines.append(f"  {control_id}_eval_{idx}: {str(label).strip()}")
    for label in ev_obj.get("fail_if") or []:
        idx += 1
        lines.append(f"  {control_id}_eval_f{idx}: [FAIL IF] {str(label).strip()}")
    for label in ev_obj.get("cross_checks") or []:
        idx += 1
        lines.append(f"  {control_id}_eval_c{idx}: [CROSS-CHECK] {str(label).strip()}")
    return "\n".join(lines)


def _eval_criteria_pass_if_only(ev_raw: str | dict | Any | None) -> list[tuple[str, str]]:
    """Return [(id, label), ...] for UI display: only pass_if from evaluation_criteria JSON."""
    ev_obj = _parse_evaluation_criteria_structured(ev_raw)
    if not ev_obj:
        return _parse_numbered_criteria(ev_raw)
    return [(str(i), str(x).strip()) for i, x in enumerate(ev_obj.get("pass_if") or [], 1) if str(x).strip()]


def _format_criteria_as_numbered_list(parsed: list[tuple[str, str]]) -> str:
    """Turn [(id, label), ...] into '1. label\\n2. label\\n...' for prompt."""
    return "\n".join(f"{id_}. {label}" for id_, label in parsed)


def _format_criteria_with_control_ids(control_id: str, parsed: list[tuple[str, str]], prefix: str = "") -> str:
    """Format criteria as 'control_id_prefix_N: label' so the LLM uses those exact IDs in its response.
    prefix should be 'suf_' for sufficiency or 'eval_' for evaluation to avoid ID collisions."""
    return "\n".join(f"  {control_id}_{prefix}{id_}: {label}" for id_, label in parsed)


def _matrix_row_field(row: Any, name: str) -> Any:
    """Read a field from either a dict row or an ORM model (e.g. EvidenceSufficiencyMatrix)."""
    if isinstance(row, dict):
        return row.get(name)
    return getattr(row, name, None)


def _build_prompt(
    item: Any,
    mappings: list[Any],
    matrix_rows: list[Any] | None = None,
    submission_context: str | None = None,
    previous_evaluation: dict | None = None,
) -> str:
    """Build the evaluation prompt from evidence_sufficiency_matrix rows when present, else from CEI fields.
    When matrix_rows is set, all controls' sufficiency_criteria and evaluation_criteria are included
    as readable numbered lists. submission_context (e.g. form data) is appended for the AI."""
    controls_text = "\n".join(
        f"- {m.control_id}: {getattr(m, 'sufficiency_requirement', None) or 'General compliance'}"
        for m in mappings
    )
    if matrix_rows:
        sufficiency_parts = []
        evaluation_parts = []
        for row in matrix_rows:
            cid = _matrix_row_field(row, "control_id")
            cname = _matrix_row_field(row, "control_name") or cid
            suf_raw = _matrix_row_field(row, "sufficiency_criteria")
            ev_raw = _matrix_row_field(row, "evaluation_criteria")
            suf_parsed = _parse_numbered_criteria(suf_raw)
            ev_structured = _parse_evaluation_criteria_structured(ev_raw)
            if suf_parsed:
                sufficiency_parts.append(
                    f"--- Control {cid} ({cname}) ---\n" + _format_criteria_with_control_ids(cid, suf_parsed, prefix="suf_")
                )
            if ev_structured:
                evaluation_parts.append(
                    f"--- Control {cid} ({cname}) ---\n" + _format_evaluation_structured_for_prompt(cid, ev_structured)
                )
            elif ev_raw:
                ev_parsed = _parse_numbered_criteria(ev_raw)
                if ev_parsed:
                    evaluation_parts.append(
                        f"--- Control {cid} ({cname}) ---\n" + _format_criteria_with_control_ids(cid, ev_parsed, prefix="eval_")
                    )
        sufficiency_definition = "\n\n".join(sufficiency_parts) if sufficiency_parts else "N/A"
        evaluation_criteria = "\n\n".join(evaluation_parts) if evaluation_parts else "N/A"
        evidence_description = (
            getattr(item, "evidence_description", None) or "See per-control sufficiency and evaluation criteria below."
        )
    else:
        evidence_description = getattr(item, "evidence_description", None) or "N/A"
        sufficiency_definition = getattr(item, "sufficiency_definition", None) or "N/A"
        evaluation_criteria = getattr(item, "evaluation_criteria", None) or "N/A"
    template = _load_prompt_template()
    form_fields_section = (
        submission_context.strip() if (submission_context and submission_context.strip()) else "None provided."
    )
    out = template.format(
        evidence_item_id=item.id,
        evidence_item_name=item.name,
        evidence_description=evidence_description,
        sufficiency_definition=sufficiency_definition,
        evaluation_criteria=evaluation_criteria,
        mapped_controls=controls_text or "None",
    )
    out = out.replace("<<<FORM_FIELDS_SECTION>>>", form_fields_section)
    out += (
        "\n\n## CRITICAL RESPONSE RULES:\n"
        "- IDs starting with 'suf_' (e.g. 1.1_suf_1) are SUFFICIENCY criteria. Return them ONLY in the 'sufficiency_results' array.\n"
        "- IDs starting with 'eval_' (e.g. 1.1_eval_1) are EVALUATION criteria. Return them ONLY in the 'criteria' array.\n"
        "- NEVER duplicate: each criterion ID must appear in exactly ONE array.\n"
        "- Use the EXACT IDs provided (including the suf_/eval_ prefix).\n"
        "- For met=false criteria, keep description concise and user-friendly: max 2 short points, prefer 1.\n"
        "- Description style: what is missing + what to add/verify; avoid long paragraphs.\n"
    )

    if previous_evaluation:
        prev_lines = []
        met_ids = []
        for section_key in ("sufficiency_results", "criteria"):
            for c in previous_evaluation.get(section_key, []):
                status = "MET" if c.get("met") else "NOT MET"
                cid = c.get("id") or ""
                if c.get("met"):
                    met_ids.append(cid)
                desc = c.get("description") or ""
                label = c.get("label", c.get("id", ""))
                line = f"  - [{status}] {label}"
                if desc:
                    line += f" — {desc}"
                prev_lines.append(line)
        if prev_lines:
            out += (
                "\n\n## Previous evaluation results (user may have edited these; consider changes as user corrections):\n"
                + "\n".join(prev_lines)
            )
            if met_ids:
                out += (
                    "\n\n## CRITICAL — Regenerate rule:\n"
                    "Criteria marked as [MET] above (user has confirmed these are satisfied) MUST remain met=true in your response. "
                    "Do NOT return them as met=false. IDs that must stay MET: "
                    + ", ".join(met_ids)
                    + "\n"
                )

    return out


def _parse_ai_response(text: str) -> dict:
    """Extract JSON from the model response, tolerating markdown fences."""
    cleaned = text.strip()
    if cleaned.startswith("```"):
        first_nl = cleaned.index("\n")
        cleaned = cleaned[first_nl + 1 :]
    if cleaned.endswith("```"):
        cleaned = cleaned[: cleaned.rfind("```")]
    return json.loads(cleaned.strip())


def _preview_text(text: str, max_chars: int = _AI_LOG_PREVIEW_CHARS) -> str:
    """Return response preview for terminal logs without flooding output."""
    if len(text) <= max_chars:
        return text
    return f"{text[:max_chars]}\n...[truncated {len(text) - max_chars} chars]"


def _log_ai_response(label: str, raw_text: str) -> None:
    """Print raw response to CMD, then optional pretty JSON summary for debugging."""
    _ai_raw_response_console(label, raw_text)
    try:
        parsed = _parse_ai_response(raw_text)
        pretty = json.dumps(parsed, indent=2, ensure_ascii=False, default=str)
        _ai_emit("INFO", label, "Parsed JSON (pretty preview):\n" + _preview_text(pretty, _AI_LOG_PREVIEW_CHARS))
    except Exception:
        _ai_emit("WARN", label, "Model output is not valid JSON after fence strip; see RAW block above.")


def evaluate_evidence(
    file_parts: list[Part],
    evidence_item: Any,
    control_mappings: list[Any],
    matrix_rows: list[Any] | None = None,
    submission_context: str | None = None,
    previous_evaluation: dict | None = None,
) -> dict:
    """Send files + prompt to Vertex AI and return the parsed JSON result.
    When matrix_rows is provided (from evidence_sufficiency_matrix), prompt is built from them.
    submission_context: optional (e.g. A5 declared architecture and form data) appended to prompt."""
    try:
        model = _get_model()
    except Exception as init_err:
        _ai_emit("ERROR", "evaluate_evidence", f"Vertex AI init failed: {init_err}")
        logger.exception("Vertex AI init failed")
        raise ValueError(
            "Vertex AI is not available. Check GOOGLE_CLOUD_PROJECT and "
            "Application Default Credentials (gcloud auth application-default login)."
        ) from init_err

    prompt_text = _build_prompt(
        evidence_item, control_mappings, matrix_rows=matrix_rows,
        submission_context=submission_context, previous_evaluation=previous_evaluation,
    )
    _ai_emit(
        "INFO",
        "evaluate_evidence",
        "start "
        f"item={getattr(evidence_item, 'id', 'unknown')} "
        f"controls={len(control_mappings)} files={len(file_parts)} "
        f"model={settings.VERTEX_AI_MODEL} prompt_chars={len(prompt_text)}",
    )
    contents = [Part.from_text(prompt_text)] + file_parts

    max_retries = 3
    last_err: Exception | None = None
    for attempt in range(max_retries + 1):
        try:
            response = model.generate_content(contents)
            _ai_emit(
                "INFO",
                "evaluate_evidence",
                f"vertex_ok attempt={attempt + 1}/{max_retries + 1}",
            )
            break  # success
        except Exception as api_err:
            last_err = api_err
            err_msg = str(api_err)
            err_type = type(api_err).__name__

            is_rate_limit = "RESOURCE_EXHAUSTED" in err_type or "ResourceExhausted" in err_type or "429" in err_msg or "Resource exhausted" in err_msg
            if is_rate_limit and attempt < max_retries:
                wait = 2 ** attempt * 5  # 5s, 10s, 20s
                _ai_emit(
                    "WARN",
                    "evaluate_evidence",
                    f"rate_limit retry in {wait}s (attempt {attempt + 1}/{max_retries + 1}) err={err_msg[:200]}",
                )
                logger.warning("Vertex AI 429 (attempt %d/%d), retrying in %ds…", attempt + 1, max_retries + 1, wait)
                time.sleep(wait)
                continue

            if is_rate_limit:
                _ai_emit("WARN", "evaluate_evidence", f"rate_limit exhausted after {attempt + 1} attempts: {err_msg}")
                logger.warning("Vertex AI rate limit (429) after %d attempts: %s", attempt + 1, err_msg)
                raise ValueError(
                    "Vertex AI rate limit (429). The AI service is temporarily overloaded. Please try again in a minute."
                ) from api_err

            _ai_emit("ERROR", "evaluate_evidence", f"generate_content failed: {err_type}: {err_msg}")
            logger.exception("Vertex AI generate_content failed")
            if "PERMISSION_DENIED" in err_msg or "PermissionDenied" in err_type or "403" in err_msg:
                raise ValueError(
                    "Vertex AI permission denied (403). To fix: (1) Enable Vertex AI API in Google Cloud Console "
                    "(https://console.cloud.google.com/apis/library/aiplatform.googleapis.com?project=YOUR_PROJECT). "
                    "(2) Grant the service account or user the role 'Vertex AI User' (roles/aiplatform.user) so it has "
                    "aiplatform.endpoints.predict. (3) If the model does not exist in your region, set VERTEX_AI_MODEL "
                    "to a supported model. Original error: " + err_msg
                ) from api_err
            raise ValueError(f"Vertex AI API error: {api_err}") from api_err
    else:
        raise ValueError(f"Vertex AI failed after {max_retries + 1} attempts: {last_err}")

    text = _get_response_text(response)
    if not (text and text.strip()):
        raise ValueError(
            "Vertex AI returned no text (response may have been blocked or empty). "
            "Try different evidence files or check model safety settings."
        )

    _log_ai_response("evaluate_evidence", text)
    try:
        parsed = _parse_ai_response(text)
        _ai_emit("INFO", "evaluate_evidence", "parse_ok returning dict to caller")
        return parsed
    except (json.JSONDecodeError, ValueError) as exc:
        _ai_emit("ERROR", "evaluate_evidence", f"json_parse_failed: {exc}")
        logger.error("AI response parse error: %s\nRaw: %s", exc, _preview_text(text, 2000))
        raise ValueError(f"AI returned invalid JSON: {exc}") from exc


def _get_response_text(response: Any) -> str:
    """Extract generated text from Vertex AI GenerateContentResponse."""
    if hasattr(response, "text") and response.text is not None:
        return response.text
    if hasattr(response, "candidates") and response.candidates:
        for c in response.candidates:
            if c.content and c.content.parts:
                for p in c.content.parts:
                    if getattr(p, "text", None):
                        return p.text
    return ""


# ---------------------------------------------------------------------------
# Report generation
# ---------------------------------------------------------------------------

DOMAIN_NAMES = {
    "A": "Network & Architecture",
    "B": "System Hardening & Config",
    "C": "Access Management",
    "D": "Vulnerability & Patch Mgmt",
    "E": "Monitoring & Detection",
    "F": "Third-Party & Outsourcing",
    "G": "Physical Security",
    "H": "Policies & Governance",
}

_SECTION_INSTRUCTIONS: dict[str, str] = {
    "executive_summary": (
        "Write a 1–2 page Executive Summary for the SWIFT CSP assessment report.\n"
        "Include:\n"
        "- Purpose of the assessment\n"
        "- Overall compliance status (use the numbers provided, do NOT recalculate)\n"
        "- A compliance status table (Total / Mandatory / Advisory / Compliant / Partially / Non-Compliant / N/A)\n"
        "- Key observations (strengths and weaknesses)\n"
        "- High-level risk posture\n"
        "- Overall conclusion statement: 'Based on the procedures performed, [Bank] is assessed as [status] with the applicable SWIFT CSCF controls.'"
    ),
    "scope_methodology": (
        "Write the Scope of Assessment and Methodology section.\n"
        "Include:\n"
        "- Architecture type and what it means\n"
        "- In-scope systems (Messaging Interface, Communication Interface, SwiftNet Link, HSM, GUI)\n"
        "- Assessment period\n"
        "- Methodology: evidence review, L1 (completeness) → L2 (quality) → L3 (assessment) review workflow\n"
        "- Testing approach: documentation inspection, configuration validation, sample-based testing\n"
        "- Do NOT list actual evidence items; keep it at methodology level."
    ),
    "gap_analysis": (
        "Write a consolidated Gap Analysis section.\n"
        "Include:\n"
        "- A summary table: Control ID | Control Name | Status | Risk Level | Target Remediation\n"
        "- Only include controls that are NOT fully compliant\n"
        "- For each gap, write a brief observation and recommendation\n"
        "- Systemic patterns if any\n"
        "- Risk distribution summary (high / medium / low counts)"
    ),
    "attestation": (
        "Write the Final Attestation section.\n"
        "Include:\n"
        "- Formal attestation statement in SWIFT IAF format\n"
        "- Approved by, role, date, MFA verification status\n"
        "- If not yet attested, state that attestation is pending."
    ),
    "evidence_index": (
        "Generate an Evidence Index table listing all evidence items.\n"
        "Columns: Evidence Item ID | Status | Overall Met | Domain\n"
        "List every item from the data provided."
    ),
}


def _get_section_instructions(section_key: str) -> str:
    if section_key in _SECTION_INSTRUCTIONS:
        return _SECTION_INSTRUCTIONS[section_key]

    if section_key.startswith("domain_") and len(section_key) == 8:
        domain_letter = section_key[-1].upper()
        domain_name = DOMAIN_NAMES.get(domain_letter, f"Domain {domain_letter}")
        return (
            f"Write the detailed assessment section for Domain {domain_letter} — {domain_name}.\n"
            "For EACH control in this domain, write:\n"
            "### Control [ID] — [Name]\n"
            "- **Control Objective:** (from the data)\n"
            "- **Control Type:** Mandatory / Advisory\n"
            "- **Evidence Reviewed:** list the evidence item IDs and their status\n"
            "- **Testing Performed:** describe what was reviewed based on the evidence status and evaluation\n"
            "- **Assessment Result:** Compliant / Partially Compliant / Non-Compliant (use the compliance_status from data)\n"
            "- **Risk Rating:** High / Medium / Low\n"
            "- **Observation:** describe what was found based on eval_summary and failed_criteria\n"
            "- **Recommendation:** if non-compliant or partially compliant, what should be done (use remediation data)\n"
            "\nIf a control has no evidence items, note that evidence was not available."
        )

    return "Write the content for this section based on the provided data."


def _slice_snapshot_for_section(snapshot: dict, section_key: str) -> dict:
    """Return only the data slice relevant to a section to keep prompts focused."""
    meta = snapshot.get("metadata", {})
    summary = snapshot.get("overall_summary", {})
    controls = snapshot.get("controls", [])
    risk = snapshot.get("risk_summary", {})
    att = snapshot.get("attestation", {})

    if section_key == "executive_summary":
        return {"metadata": meta, "overall_summary": summary, "risk_summary": risk, "attestation": att}

    if section_key == "scope_methodology":
        return {"metadata": meta, "overall_summary": {"total_controls": summary.get("total_controls"),
                "mandatory_controls": summary.get("mandatory_controls"),
                "advisory_controls": summary.get("advisory_controls")}}

    if section_key.startswith("domain_") and len(section_key) == 8:
        domain_letter = section_key[-1].upper()
        domain_controls = [c for c in controls if c.get("domain", "").upper() == domain_letter]
        return {"metadata": meta, "controls": domain_controls}

    if section_key == "gap_analysis":
        non_compliant = [c for c in controls if c.get("compliance_status") != "compliant"]
        return {"metadata": meta, "risk_summary": risk, "controls": non_compliant}

    if section_key == "attestation":
        return {"metadata": meta, "attestation": att}

    if section_key == "evidence_index":
        items = []
        for c in controls:
            for ev in c.get("evidence_items", []):
                items.append({
                    "evidence_item_id": ev["evidence_item_id"],
                    "status": ev["status"],
                    "overall_met": ev.get("overall_met"),
                    "domain": c.get("domain", "?"),
                })
        seen = set()
        unique = []
        for item in items:
            if item["evidence_item_id"] not in seen:
                seen.add(item["evidence_item_id"])
                unique.append(item)
        return {"evidence_items": sorted(unique, key=lambda x: x["evidence_item_id"])}

    return snapshot


_SECTION_DISPLAY_NAMES: dict[str, str] = {
    "executive_summary": "Executive Summary",
    "scope_methodology": "Scope & Methodology",
    "gap_analysis": "Gap Analysis",
    "attestation": "Final Attestation",
    "evidence_index": "Evidence Index",
    "glossary": "Glossary",
}


def _section_display_name(section_key: str) -> str:
    if section_key in _SECTION_DISPLAY_NAMES:
        return _SECTION_DISPLAY_NAMES[section_key]
    if section_key.startswith("domain_") and len(section_key) == 8:
        letter = section_key[-1].upper()
        return f"Domain {letter} — {DOMAIN_NAMES.get(letter, letter)}"
    return section_key


def _load_report_prompt_template() -> str:
    if not _REPORT_PROMPT_FILE.is_file():
        raise FileNotFoundError(f"Report prompt file not found: {_REPORT_PROMPT_FILE}")
    return _REPORT_PROMPT_FILE.read_text(encoding="utf-8")


def _load_aws_autofill_prompt_template() -> str:
    if not _AWS_AUTOFILL_PROMPT_FILE.is_file():
        raise FileNotFoundError(f"AWS autofill prompt file not found: {_AWS_AUTOFILL_PROMPT_FILE}")
    return _AWS_AUTOFILL_PROMPT_FILE.read_text(encoding="utf-8")


def _load_gcp_autofill_prompt_template() -> str:
    if not _GCP_AUTOFILL_PROMPT_FILE.is_file():
        raise FileNotFoundError(f"GCP autofill prompt file not found: {_GCP_AUTOFILL_PROMPT_FILE}")
    return _GCP_AUTOFILL_PROMPT_FILE.read_text(encoding="utf-8")


def generate_report_section(snapshot: dict, section_key: str) -> str:
    """Generate a single report section via Vertex AI.

    Returns markdown text for the section content.
    """
    if section_key == "glossary":
        return _static_glossary()

    meta = snapshot.get("metadata", {})
    section_data = _slice_snapshot_for_section(snapshot, section_key)
    instructions = _get_section_instructions(section_key)
    section_name = _section_display_name(section_key)

    template = _load_report_prompt_template()
    prompt = template.format(
        bank_name=meta.get("bank_name", "Unknown"),
        bic_code=meta.get("bic_code", "N/A"),
        assessment_year=meta.get("assessment_year", "N/A"),
        architecture_type=meta.get("architecture_type", "Unknown"),
        assessment_period=meta.get("assessment_period", "N/A"),
        cscf_version=meta.get("cscf_version", "2025v"),
        section_key=section_key,
        section_instructions=instructions,
        section_data=json.dumps(section_data, indent=2, default=str),
        section_name=section_name,
    )
    _ai_emit(
        "INFO",
        "generate_report_section",
        f"start section={section_key} name={section_name} model={settings.VERTEX_AI_MODEL} prompt_chars={len(prompt)}",
    )

    model = _get_model()

    max_retries = 5
    for attempt in range(max_retries + 1):
        try:
            response = model.generate_content([Part.from_text(prompt)])
            _ai_emit(
                "INFO",
                "generate_report_section",
                f"vertex_ok section={section_key} attempt={attempt + 1}/{max_retries + 1}",
            )
            break
        except Exception as api_err:
            err_msg = str(api_err)
            is_rate_limit = "429" in err_msg or "RESOURCE_EXHAUSTED" in err_msg
            if is_rate_limit and attempt < max_retries:
                wait = min(5 * (2 ** attempt), 90)
                _ai_emit(
                    "WARN",
                    "generate_report_section",
                    f"section={section_key} rate_limit retry in {wait}s attempt={attempt + 1}/{max_retries + 1}",
                )
                logger.warning(
                    "Report gen 429 (attempt %d/%d), retrying in %ds…",
                    attempt + 1, max_retries + 1, wait,
                )
                time.sleep(wait)
                continue
            if is_rate_limit:
                _ai_emit("ERROR", "generate_report_section", f"section={section_key} rate_limit_gave_up: {err_msg}")
                raise ValueError(
                    f"Vertex AI rate limit (429). Please try again in a few minutes. "
                    f"Section: {section_key}. Error: {api_err}"
                ) from api_err
            _ai_emit("ERROR", "generate_report_section", f"section={section_key} vertex_error: {err_msg}")
            raise ValueError(f"Vertex AI error (section={section_key}): {api_err}") from api_err

    text = _get_response_text(response)
    if not text or not text.strip():
        _ai_emit("WARN", "generate_report_section", f"section={section_key} empty_response")
        return f"*Content generation returned empty for section: {section_name}. Please regenerate.*"

    _ai_raw_response_console(f"generate_report_section:{section_key}", text)
    cleaned = text.strip()
    if cleaned.startswith("```"):
        first_nl = cleaned.index("\n") if "\n" in cleaned else len(cleaned)
        cleaned = cleaned[first_nl + 1:]
    if cleaned.endswith("```"):
        cleaned = cleaned[:cleaned.rfind("```")]

    out = cleaned.strip()
    _ai_emit("INFO", "generate_report_section", f"section={section_key} done markdown_chars={len(out)}")
    return out


_DEFAULT_SUGGESTION_GAP = (
    "The cloud snapshots for this cycle do not contain a clear signal for this field, or the relevant "
    "collector has not been run yet. Run collection for this CSCF item or answer manually."
)


def _aws_suggestion_value_is_usable(question_type: str, raw: str | None) -> bool:
    """
    True when the model returned something we treat as a real answer (not empty / not [] spreadsheet).
    ``[]`` must be false so ``gaps`` from the LLM are still returned to the UI.
    """
    qt = (question_type or "").strip().lower()
    s = "" if raw is None else str(raw).strip()
    if not s:
        return False
    if qt == "spreadsheet":
        if s in ("[]", "{}", "null", "undefined"):
            return False
        try:
            rows = json.loads(s)
        except (json.JSONDecodeError, TypeError):
            return False
        if not isinstance(rows, list) or len(rows) == 0:
            return False
        return any(
            isinstance(row, dict) and any(str(v if v is not None else "").strip() for v in row.values())
            for row in rows
        )
    if qt == "checkbox":
        return s in ("true", "false")
    return True


def log_suggest_from_aws_skip(reason: str, detail: str = "") -> None:
    """Call from HTTP layer when suggest-from-aws returns without calling Vertex (visible in uvicorn terminal)."""
    log_suggest_from_cloud_skip(reason, detail, cloud_provider="aws")


def log_suggest_from_cloud_skip(
    reason: str, detail: str = "", *, cloud_provider: str | None = None
) -> None:
    """Call from HTTP layer when cloud autofill returns without calling Vertex."""
    msg = f"skip_llm reason={reason}"
    if detail:
        msg = f"{msg} | {detail}"
    if cloud_provider:
        msg = f"{msg} | provider={cloud_provider}"
    _ai_emit("WARN", "suggest_cloud_autofill", msg)


def suggest_answers_from_aws_evidence(
    *,
    evidence_item_id: str,
    questions: list[dict],
    aws_evidence_bundle: list[dict],
    collector_warning: str | None = None,
) -> dict[str, Any]:
    return suggest_answers_from_cloud_evidence(
        evidence_item_id=evidence_item_id,
        questions=questions,
        evidence_bundle=aws_evidence_bundle,
        cloud_provider="aws",
        collector_warning=collector_warning,
    )


def suggest_answers_from_gcp_evidence(
    *,
    evidence_item_id: str,
    questions: list[dict],
    gcp_evidence_bundle: list[dict],
    collector_warning: str | None = None,
) -> dict[str, Any]:
    return suggest_answers_from_cloud_evidence(
        evidence_item_id=evidence_item_id,
        questions=questions,
        evidence_bundle=gcp_evidence_bundle,
        cloud_provider="gcp",
        collector_warning=collector_warning,
    )


def suggest_answers_from_azure_evidence(
    *,
    evidence_item_id: str,
    questions: list[dict],
    azure_evidence_bundle: list[dict],
    collector_warning: str | None = None,
) -> dict[str, Any]:
    """Same pipeline as GCP autofill (JSON evidence bundle); uses GCP-oriented prompt for Azure Resource Graph payloads."""
    return suggest_answers_from_cloud_evidence(
        evidence_item_id=evidence_item_id,
        questions=questions,
        evidence_bundle=azure_evidence_bundle,
        cloud_provider="azure",
        collector_warning=collector_warning,
    )


_vertex_thread_pool = ThreadPoolExecutor(max_workers=4, thread_name_prefix="vertex")


def _generate_with_timeout(model: GenerativeModel, parts: list, gen_config, timeout_s: int):
    """Run model.generate_content in a thread with a hard timeout so the request never hangs."""
    future = _vertex_thread_pool.submit(model.generate_content, parts, generation_config=gen_config)
    try:
        return future.result(timeout=timeout_s)
    except FuturesTimeoutError:
        future.cancel()
        raise TimeoutError(
            f"Vertex AI generate_content did not respond in {timeout_s}s. "
            "Check your network, credentials, and Google Cloud project quotas."
        )


def suggest_answers_from_cloud_evidence(
    *,
    evidence_item_id: str,
    questions: list[dict],
    evidence_bundle: list[dict],
    cloud_provider: str = "aws",
    collector_warning: str | None = None,
) -> dict[str, Any]:
    """
    Use Vertex AI to map cloud collector JSON into form question_key -> string answers.
    Provider selects ``aws_autofill`` vs ``gcp_autofill`` prompt templates (Azure uses the GCP template).
    `questions`: list built by the evidence router (includes mapping columns per provider).

    Returns a dict: ``{"suggestions": {question_key: str, ...}, "gaps": {question_key: str, ...}}``.
    ``gaps`` explains empty answers (missing or unmappable evidence).
    """
    if not questions:
        return {"suggestions": {}, "gaps": {}}
    if not evidence_bundle:
        return {"suggestions": {}, "gaps": {}}
    provider = (cloud_provider or "aws").strip().lower()
    if provider not in {"aws", "gcp", "azure"}:
        raise ValueError("Unsupported cloud provider. Use 'aws', 'gcp', or 'azure'.")

    log_op = {
        "aws": "suggest_answers_from_aws_evidence",
        "gcp": "suggest_answers_from_gcp_evidence",
        "azure": "suggest_answers_from_azure_evidence",
    }[provider]
    try:
        model = _get_model()
    except Exception as init_err:
        _ai_emit("ERROR", log_op, f"init_failed: {init_err}")
        logger.exception("Vertex AI init failed (suggest-from-cloud)")
        raise ValueError(
            "Vertex AI is not available. Set GOOGLE_CLOUD_PROJECT and credentials."
        ) from init_err

    # Compact JSON: indented dumps balloon token count and slow Vertex.
    _jc = lambda obj: json.dumps(obj, default=str, separators=(",", ":"))
    q_json = _jc(questions)
    ev_json = _jc(evidence_bundle)
    prompt_template = (
        _load_aws_autofill_prompt_template()
        if provider == "aws"
        else _load_gcp_autofill_prompt_template()
    )
    # AWS template uses {aws_evidence_json}; GCP/Azure use {cloud_evidence_json}. Python str.format()
    # rejects unexpected keyword arguments — passing both caused TypeError and HTTP 500 on suggest-from-gcp.
    if provider == "aws":
        prompt = prompt_template.format(
            evidence_item_id=evidence_item_id,
            aws_evidence_json=ev_json,
            questions_json=q_json,
        )
    else:
        prompt = prompt_template.format(
            evidence_item_id=evidence_item_id,
            cloud_evidence_json=ev_json,
            questions_json=q_json,
        )
    prompt += (
        "\n\n## RUNTIME — RESPONSE RULES (do not echo this section)\n"
        "- Output one JSON object only; keys exactly `answers` and `gaps`.\n"
        "- Every listed question_key must appear in `answers` (string values).\n"
        "- Empty strings need a one-line reason in `gaps` under the same key.\n"
    )
    cw = (collector_warning or "").strip()
    if cw:
        prompt += (
            "\n\n## COLLECTOR STATUS (runtime — do not echo this section)\n"
            f"Summary: {cw}\n"
            "The evidence JSON may contain only collector error payloads. Do not invent facts. "
            "If errors prevent mapping, leave answers empty and explain in `gaps` (API disabled, IAM, project config, etc.).\n"
        )
    _ai_emit(
        "INFO",
        log_op,
        f"start provider={provider} item={evidence_item_id} questions={len(questions)} evidence_chunks={len(evidence_bundle)} "
        f"model={settings.VERTEX_AI_MODEL} prompt_chars={len(prompt)} "
        "| console: this Python/uvicorn process (not Next.js npm run dev)",
    )
    _ai_emit(
        "INFO",
        log_op,
        "prompt_preview (truncated):\n" + _preview_text(prompt, min(_AI_LOG_PREVIEW_CHARS, 12_000)),
    )

    suggest_gen_cfg = GenerationConfig(
        temperature=0.2,
        max_output_tokens=settings.LLM_SUGGEST_MAX_OUTPUT_TOKENS,
    )
    max_retries = 3
    last_err: Exception | None = None
    response = None
    for attempt in range(max_retries + 1):
        try:
            timeout_s = settings.VERTEX_GENERATE_TIMEOUT_S
            _ai_emit(
                "INFO",
                log_op,
                f"vertex_generate_content_start timeout={timeout_s}s (blocks until Google API returns — watch elapsed_s on next line)",
            )
            t_vertex = time.monotonic()
            response = _generate_with_timeout(
                model,
                [Part.from_text(prompt)],
                suggest_gen_cfg,
                timeout_s,
            )
            elapsed = time.monotonic() - t_vertex
            _ai_emit(
                "INFO",
                log_op,
                f"vertex_ok attempt={attempt + 1}/{max_retries + 1} elapsed_s={elapsed:.1f}",
            )
            break
        except Exception as api_err:
            last_err = api_err
            err_msg = str(api_err)
            err_type = type(api_err).__name__
            is_rate_limit = (
                "RESOURCE_EXHAUSTED" in err_type
                or "ResourceExhausted" in err_type
                or "429" in err_msg
                or "Resource exhausted" in err_msg
            )
            if is_rate_limit and attempt < max_retries:
                wait = min(2**attempt * 5, 60)
                _ai_emit(
                    "WARN",
                    log_op,
                    f"rate_limit retry in {wait}s attempt={attempt + 1}/{max_retries + 1} err={err_msg[:240]}",
                )
                time.sleep(wait)
                continue
            if is_rate_limit:
                _ai_emit(
                    "ERROR",
                    log_op,
                    f"rate_limit exhausted: {err_msg[:400]}",
                )
                raise ValueError(
                    "Vertex AI rate limit (429). Try again in a minute."
                ) from api_err
            _ai_emit(
                "ERROR",
                log_op,
                f"generate_content failed: {err_type}: {err_msg[:500]}",
            )
            raise ValueError(
                f"Vertex AI API error (suggest-from-{provider}): {api_err}"
            ) from api_err
    else:
        raise ValueError(f"Vertex AI failed after {max_retries + 1} attempts: {last_err}")

    text = _get_response_text(response)
    if not (text and text.strip()):
        _ai_emit("ERROR", log_op, "empty_model_text")
        raise ValueError(f"Vertex AI returned empty text for suggest-from-{provider}.")
    _log_ai_response(log_op, text)

    try:
        parsed = _parse_ai_response(text)
    except (json.JSONDecodeError, ValueError) as e:
        _ai_emit("ERROR", log_op, f"json_parse_failed: {e}")
        raise ValueError(f"Model returned invalid JSON: {e}") from e
    if not isinstance(parsed, dict):
        _ai_emit("ERROR", log_op, "response_not_json_object")
        raise ValueError("Model response was not a JSON object.")
    _ai_emit("INFO", log_op, "parse_ok")

    allowed = {q["question_key"] for q in questions if q.get("question_key")}
    gaps_raw: dict[str, str] = {}
    answers_block: dict[str, Any] = {}
    if "answers" in parsed and isinstance(parsed.get("answers"), dict):
        answers_block = dict(parsed["answers"])
        g = parsed.get("gaps")
        if isinstance(g, dict):
            gaps_raw = {str(k): str(v).strip() for k, v in g.items() if v is not None and str(v).strip()}
    else:
        # Legacy: flat question_key -> value
        answers_block = {k: v for k, v in parsed.items() if k in allowed}

    out: dict[str, str] = {}
    for k in allowed:
        v = answers_block.get(k)
        out[str(k)] = "" if v is None else str(v).strip()

    key_to_type = {str(q["question_key"]): str(q.get("question_type") or "").lower() for q in questions if q.get("question_key")}

    gaps: dict[str, str] = {}
    for k in allowed:
        ks = str(k)
        qtype = key_to_type.get(ks, "")
        raw = out[ks]
        if _aws_suggestion_value_is_usable(qtype, raw):
            continue
        reason = gaps_raw.get(k) or gaps_raw.get(ks)
        gaps[ks] = (reason.strip() if reason else _DEFAULT_SUGGESTION_GAP)

    filled = sum(
        1 for k in allowed if _aws_suggestion_value_is_usable(key_to_type.get(str(k), ""), out[str(k)])
    )
    answered_keys = [
        k for k in sorted(allowed) if _aws_suggestion_value_is_usable(key_to_type.get(str(k), ""), out[str(k)])
    ]
    empty_keys = [
        k for k in sorted(allowed) if not _aws_suggestion_value_is_usable(key_to_type.get(str(k), ""), out[str(k)])
    ]
    summary_obj = {
        "evidence_item_id": evidence_item_id,
        "answered_count": filled,
        "empty_count": len(empty_keys),
        "answered_keys": answered_keys,
        "empty_keys": empty_keys,
        "gaps_preview": {k: gaps.get(k, "")[:200] for k in empty_keys[:15]},
    }
    _ai_emit(
        "INFO",
        log_op,
        "result_summary:\n" + json.dumps(summary_obj, indent=2, ensure_ascii=False),
    )
    _ai_emit(
        "INFO",
        log_op,
        f"done item={evidence_item_id} answered={filled} gaps={len(gaps)}",
    )
    return {"suggestions": out, "gaps": gaps}


def _static_glossary() -> str:
    return (
        "| Term | Definition |\n"
        "|------|------------|\n"
        "| SWIFT | Society for Worldwide Interbank Financial Telecommunication |\n"
        "| CSP | Customer Security Programme |\n"
        "| CSCF | Customer Security Controls Framework |\n"
        "| IAF | Independent Assessment Framework |\n"
        "| BIC | Business Identifier Code |\n"
        "| MFA | Multi-Factor Authentication |\n"
        "| HSM | Hardware Security Module |\n"
        "| SIEM | Security Information and Event Management |\n"
        "| L1 | Level 1 Review — Completeness check |\n"
        "| L2 | Level 2 Review — Quality / technical accuracy check |\n"
        "| L3 | Level 3 Review — Independent assessment |\n"
        "| Mandatory (M) | Control that must be implemented |\n"
        "| Advisory (A) | Control that is recommended but not required |\n"
    )
