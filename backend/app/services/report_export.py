"""
Export service for SWIFT CSP audit reports.

Generates Word (.docx) and PDF documents from report sections + metadata.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Any

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Word (docx) export
# ---------------------------------------------------------------------------

def export_docx(sections: list[dict], metadata: dict) -> io.BytesIO:
    """Generate a .docx file and return as a BytesIO buffer."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # ── Cover page ──────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SWIFT Customer Security Programme (CSP)")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x75)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Independent Assessment Report")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()

    info_items = [
        ("Organisation", metadata.get("bank_name", "N/A")),
        ("BIC Code", metadata.get("bic_code", "N/A")),
        ("Assessment Year", str(metadata.get("assessment_year", "N/A"))),
        ("Architecture Type", metadata.get("architecture_type", "N/A")),
        ("Assessment Period", metadata.get("assessment_period", "N/A")),
        ("CSCF Version", metadata.get("cscf_version", "N/A")),
    ]

    table = doc.add_table(rows=len(info_items), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(info_items):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)

    doc.add_page_break()

    # ── Table of contents placeholder ───────────────────────
    toc_heading = doc.add_heading("Table of Contents", level=1)
    for idx, sec in enumerate(sections, 1):
        p = doc.add_paragraph(f"{idx}. {sec.get('name', 'Untitled')}")
        p.style = doc.styles["List Number"]

    doc.add_page_break()

    # ── Sections ────────────────────────────────────────────
    for sec in sections:
        name = sec.get("name", "Untitled")
        content = sec.get("content", "")
        doc.add_heading(name, level=1)

        if not content:
            doc.add_paragraph("Content not yet generated.", style="Intense Quote")
            continue

        _render_markdown_to_docx(doc, content)
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _render_markdown_to_docx(doc: Any, markdown_text: str) -> None:
    """Best-effort markdown-to-docx renderer for headings, tables, lists, bold, italic."""
    from docx.shared import Pt

    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # Table (| ... | ... |)
        if line.strip().startswith("|") and "|" in line[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                stripped = lines[i].strip()
                if stripped.replace("|", "").replace("-", "").replace(" ", "").replace(":", "") == "":
                    i += 1
                    continue
                table_lines.append(stripped)
                i += 1

            if table_lines:
                _render_table(doc, table_lines)
            continue

        # Bullet list
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_text(p, text)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\s*\d+\.\s+(.+)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_rich_text(p, m.group(1))
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_rich_text(p, line)
        i += 1


def _add_rich_text(paragraph: Any, text: str) -> None:
    """Add text to a paragraph with basic bold/italic support."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _render_table(doc: Any, table_lines: list[str]) -> None:
    """Render markdown table lines into a docx table."""
    from docx.shared import Pt

    rows_data = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows_data.append(cells)

    if not rows_data:
        return

    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = "Light Grid Accent 1"

    for ri, row in enumerate(rows_data):
        for ci, cell_text in enumerate(row):
            if ci < num_cols:
                cell = table.rows[ri].cells[ci]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # Bold the header row
    if rows_data:
        for ci in range(min(num_cols, len(rows_data[0]))):
            for paragraph in table.rows[0].cells[ci].paragraphs:
                for run in paragraph.runs:
                    run.bold = True


# ---------------------------------------------------------------------------
# PDF export (uses HTML → weasyprint, falls back to basic reportlab)
# ---------------------------------------------------------------------------

def export_pdf(sections: list[dict], metadata: dict) -> io.BytesIO:
    """Generate a PDF and return as a BytesIO buffer.
    
    Uses a simple HTML-based approach with inline CSS, rendered to PDF.
    Falls back to a plain-text PDF if weasyprint is unavailable.
    """
    try:
        return _export_pdf_html(sections, metadata)
    except ImportError:
        logger.warning("weasyprint not installed; falling back to basic PDF")
        return _export_pdf_basic(sections, metadata)


def _export_pdf_html(sections: list[dict], metadata: dict) -> io.BytesIO:
    """Render report as HTML then convert to PDF via weasyprint."""
    import markdown
    from weasyprint import HTML

    bank = metadata.get("bank_name", "N/A")
    year = metadata.get("assessment_year", "N/A")

    section_html = ""
    for sec in sections:
        name = sec.get("name", "Untitled")
        content = sec.get("content", "")
        body = markdown.markdown(content, extensions=["tables"]) if content else "<p><em>Content not yet generated.</em></p>"
        section_html += f'<div class="section"><h1>{name}</h1>{body}</div>\n'

    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  @page {{ size: A4; margin: 2cm; @bottom-center {{ content: counter(page); }} }}
  body {{ font-family: Calibri, Arial, sans-serif; font-size: 10pt; color: #222; line-height: 1.5; }}
  h1 {{ color: #0F4C75; font-size: 16pt; border-bottom: 2px solid #0F4C75; padding-bottom: 4px; page-break-before: always; }}
  h1:first-of-type {{ page-break-before: avoid; }}
  h2 {{ color: #1B5E20; font-size: 13pt; }}
  h3 {{ color: #333; font-size: 11pt; }}
  table {{ border-collapse: collapse; width: 100%; margin: 10px 0; font-size: 9pt; }}
  th, td {{ border: 1px solid #ccc; padding: 6px 8px; text-align: left; }}
  th {{ background: #f0f4f8; font-weight: bold; }}
  .cover {{ text-align: center; padding: 120px 0 60px; page-break-after: always; }}
  .cover h1 {{ color: #0F4C75; font-size: 24pt; border: none; }}
  .cover h2 {{ color: #555; font-size: 14pt; }}
  .cover table {{ margin: 30px auto; width: 60%; }}
</style>
</head>
<body>
  <div class="cover">
    <h1>SWIFT Customer Security Programme (CSP)</h1>
    <h2>Independent Assessment Report</h2>
    <table>
      <tr><td><strong>Organisation</strong></td><td>{bank}</td></tr>
      <tr><td><strong>BIC Code</strong></td><td>{metadata.get("bic_code", "N/A")}</td></tr>
      <tr><td><strong>Assessment Year</strong></td><td>{year}</td></tr>
      <tr><td><strong>Architecture Type</strong></td><td>{metadata.get("architecture_type", "N/A")}</td></tr>
      <tr><td><strong>Assessment Period</strong></td><td>{metadata.get("assessment_period", "N/A")}</td></tr>
      <tr><td><strong>CSCF Version</strong></td><td>{metadata.get("cscf_version", "N/A")}</td></tr>
    </table>
  </div>
  {section_html}
</body>
</html>"""

    buf = io.BytesIO()
    HTML(string=html).write_pdf(buf)
    buf.seek(0)
    return buf


def _export_pdf_basic(sections: list[dict], metadata: dict) -> io.BytesIO:
    """Fallback: basic text-based PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet
    except ImportError:
        raise ImportError(
            "Neither weasyprint nor reportlab is installed. "
            "Install one of them: pip install weasyprint OR pip install reportlab"
        )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("SWIFT CSP Assessment Report", styles["Title"]))
    story.append(Spacer(1, 20))
    story.append(Paragraph(f"Organisation: {metadata.get('bank_name', 'N/A')}", styles["Normal"]))
    story.append(Paragraph(f"Year: {metadata.get('assessment_year', 'N/A')}", styles["Normal"]))
    story.append(PageBreak())

    for sec in sections:
        name = sec.get("name", "Untitled")
        content = sec.get("content", "")
        story.append(Paragraph(name, styles["Heading1"]))
        story.append(Spacer(1, 10))
        if content:
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 6))
                else:
                    safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    story.append(Paragraph(safe, styles["Normal"]))
        else:
            story.append(Paragraph("<i>Content not yet generated.</i>", styles["Normal"]))
        story.append(PageBreak())

    doc.build(story)
    buf.seek(0)
    return buf
